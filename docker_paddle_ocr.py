import os
import glob
import img2pdf
import sys

import paddle
from paddleocr import PaddleOCR
from natsort import natsorted
from docx import Document
from tqdm import tqdm
import logging

# ログ抑制
logging.getLogger("ppocr").setLevel(logging.WARNING)

# --- 設定エリア ---
# ※ Docker コンテナ内から見たパスで指定します
IMAGE_DIR = "./hokuto_scrapes/"   # 画像の保存先
OUTPUT_DIR = "./"            # 出力先
PDF_FILENAME = "_reviews.pdf"
DOCX_FILENAME = "_reviews.docx"
TXT_FILENAME = "_raw_text.txt"

CONFIDENCE_THRESHOLD = 0.6


def setup_device():
    """
    Docker 環境で GPU / CPU を自動選択するユーティリティ。
    - CUDA が使えれば GPU
    - 使えなければ CPU にフォールバック
    """
    if paddle.device.is_compiled_with_cuda():
        try:
            paddle.set_device("gpu")
            print("✅ GPUモードで動作します (paddle.set_device('gpu'))")
            return True
        except Exception as e:
            print(f"⚠️ GPU設定に失敗したため CPU にフォールバックします: {e}")

    paddle.set_device("cpu")
    print("⚠️ GPUが利用できないため CPUモードで動作します")
    return False


def main():
    # 1. 出力ディレクトリの作成
    try:
        os.makedirs(OUTPUT_DIR, exist_ok=True)
    except Exception as e:
        print(f"⚠️ ディレクトリ作成エラー: {e}")
        return

    # 2. デバイス (GPU / CPU) 設定
    use_gpu = setup_device()

    # 3. 画像リスト取得
    img_paths = natsorted(glob.glob(os.path.join(IMAGE_DIR, "*.png")))
    if not img_paths:
        print(f"エラー: '{IMAGE_DIR}' に画像が見つかりません。")
        return

    print(f"対象画像: {len(img_paths)}枚")

    # =========================================================
    # Phase 1: PDF作成（画像を1つのPDFにまとめる）
    # =========================================================
    print("\n[1/3] PDFを作成しています...")
    pdf_path = os.path.join(OUTPUT_DIR, PDF_FILENAME)

    try:
        with open(pdf_path, "wb") as f:
            f.write(img2pdf.convert(img_paths))
        print(f"  -> PDF保存完了: {pdf_path}")
    except Exception as e:
        print(f"  -> PDF作成エラー: {e}")

    # =========================================================
    # Phase 2: PaddleOCR実行（GPU / CPU）
    # =========================================================
    print("\n[2/3] PaddleOCRで解析しています...")

    ocr = PaddleOCR(
        use_angle_cls=False,  # 元のスクリプトの設定を踏襲
        lang="japan",
        use_gpu=use_gpu,      # ← ここが GPU/CPU を切り替えるポイント
        show_log=False,
    )

    full_text_data = []

    for img_path in tqdm(img_paths, desc="OCR Progress"):
        file_name = os.path.basename(img_path)

        # OCR実行
        result = ocr.ocr(img_path, cls=True)

        page_text = []

        if result and result[0]:
            for line in result[0]:
                text = line[1][0]
                score = line[1][1]

                if score >= CONFIDENCE_THRESHOLD:
                    page_text.append(text)

        full_text_data.append(
            {
                "filename": file_name,
                "content": "\n".join(page_text),
            }
        )

    # =========================================================
    # Phase 3: ファイル出力（TXT / DOCX）
    # =========================================================
    print("\n[3/3] 結果をファイルに書き出しています...")

    # テキスト出力
    txt_path = os.path.join(OUTPUT_DIR, TXT_FILENAME)
    with open(txt_path, "w", encoding="utf-8") as f:
        for page in full_text_data:
            f.write(f"--- Page: {page['filename']} ---\n")
            f.write(page["content"])
            f.write("\n\n")

    # Word 出力
    docx_path = os.path.join(OUTPUT_DIR, DOCX_FILENAME)
    doc = Document()
    doc.add_heading("Hokuto Reviews (PaddleOCR)", 0)

    for page in full_text_data:
        doc.add_heading(f"Page: {page['filename']}", level=2)
        if page["content"].strip():
            doc.add_paragraph(page["content"])
        else:
            doc.add_paragraph("(テキストなし)")

    doc.save(docx_path)

    print(f"\n✅ 完了: {OUTPUT_DIR} を確認してください")


if __name__ == "__main__":
    main()
